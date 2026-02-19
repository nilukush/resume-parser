"""
Text Extraction Service for ResuMate.

This service extracts text from various document formats including:
- PDF files (using pdfplumber with OCR fallback)
- DOCX/DOC files (using python-docx)
- TXT files (plain text)

All extraction functions are async for consistency with the async architecture.
"""

import io
from typing import Optional
from pathlib import Path

import pdfplumber
from docx import Document

from app.services.ocr_extractor import extract_text_with_ocr


class TextExtractionError(Exception):
    """Raised when text extraction fails."""

    pass


async def extract_text(file_path: str, file_content: Optional[bytes] = None) -> str:
    """
    Extract text from various document formats.

    Args:
        file_path: Path to the file (used to determine file type)
        file_content: Optional bytes content of the file

    Returns:
        Extracted text as string

    Raises:
        TextExtractionError: If extraction fails or file type is unsupported
    """
    file_extension = Path(file_path).suffix.lower()

    if file_extension == ".pdf":
        regular_text = await _extract_from_pdf(file_path, file_content)
        text = await extract_text_with_ocr(file_path, file_content, regular_text)
        return text
    elif file_extension in [".docx", ".doc"]:
        return await _extract_from_docx(file_path, file_content)
    elif file_extension == ".txt":
        return await _extract_from_txt(file_content)
    else:
        raise TextExtractionError(f"Unsupported file type: {file_extension}")


async def _extract_from_pdf(file_path: str, file_content: Optional[bytes] = None) -> str:
    """
    Extract text from PDF using pdfplumber.

    Args:
        file_path: Path to the PDF file
        file_content: Optional bytes content of the PDF

    Returns:
        Extracted text as string

    Raises:
        TextExtractionError: If PDF extraction fails
    """
    try:
        if file_content:
            # Extract from bytes
            pdf_file = io.BytesIO(file_content)
        else:
            # Extract from file path
            pdf_file = open(file_path, "rb")

        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text

        return text.strip()
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")


async def _extract_from_docx(file_path: str, file_content: Optional[bytes] = None) -> str:
    """
    Extract text from DOCX using python-docx.

    Args:
        file_path: Path to the DOCX file
        file_content: Optional bytes content of the DOCX

    Returns:
        Extracted text as string

    Raises:
        TextExtractionError: If DOCX extraction fails
    """
    try:
        if file_content:
            doc = Document(io.BytesIO(file_content))
        else:
            doc = Document(file_path)

        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        return text.strip()
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")


async def _extract_from_txt(file_content: Optional[bytes] = None) -> str:
    """
    Extract text from TXT file.

    Args:
        file_content: Bytes content of the TXT file

    Returns:
        Extracted text as string

    Raises:
        TextExtractionError: If TXT extraction fails
    """
    try:
        if file_content:
            return file_content.decode("utf-8")
        return ""
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from TXT: {str(e)}")
